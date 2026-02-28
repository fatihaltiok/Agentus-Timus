# tools/document_creator/tool.py
"""
Dokument-Erstellungs-Tools für Timus.

Unterstützte Formate:
  - PDF   (fpdf2)         → create_pdf
  - DOCX  (python-docx)  → create_docx
  - XLSX  (openpyxl)     → create_xlsx
  - TXT   (Standard)     → create_txt
  - CSV   (Standard)     → create_csv

Ausgabe: immer in results/ (wird automatisch erstellt).
"""

import asyncio
import csv
import io
import logging
import re
from datetime import datetime
from pathlib import Path

from tools.tool_registry_v2 import tool, ToolParameter as P, ToolCategory as C

log = logging.getLogger(__name__)

_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
_RESULTS_DIR  = _PROJECT_ROOT / "results"


def _safe_filename(title: str, ext: str) -> Path:
    """Erstellt einen sicheren Dateinamen aus Titel + Zeitstempel."""
    _RESULTS_DIR.mkdir(exist_ok=True)
    ts    = datetime.now().strftime("%Y%m%d_%H%M%S")
    clean = re.sub(r"[^\w\- ]", "", title)[:50].strip().replace(" ", "_")
    return _RESULTS_DIR / f"{ts}_{clean}.{ext}"


def _rel(path: Path) -> str:
    """Relativer Pfad zum Projekt-Root."""
    return str(path.relative_to(_PROJECT_ROOT))


# ── PDF ───────────────────────────────────────────────────────────

@tool(
    name="create_pdf",
    description=(
        "Erstellt ein hochwertiges PDF-Dokument aus Markdown-Inhalt via Playwright (Headless Chromium). "
        "Unterstützt vollständiges Markdown: Überschriften (#, ##, ###), Tabellen, Code-Blöcke, "
        "Listen, Fettschrift, Kursiv, Blockquotes. Saubere Seitenumbrüche, Kopf-/Fußzeile, "
        "professionelles CSS-Design. Das PDF wird in results/ gespeichert."
    ),
    parameters=[
        P("title",   "string", "Titel des Dokuments (erscheint in Kopfzeile & Metadaten)", required=True),
        P("content", "string", "Inhalt als Markdown (# = H1, ## = H2, | = Tabelle, ``` = Code)", required=True),
        P("author",  "string", "Autor (optional, erscheint in Kopfzeile)", required=False),
    ],
    capabilities=["document", "pdf"],
    category=C.DOCUMENT
)
async def create_pdf(title: str, content: str, author: str = "Timus") -> dict:
    try:
        from skills.create_clean_pdf_skill import CreateCleanPdfSkill

        out_path = _safe_filename(title, "pdf")
        skill = CreateCleanPdfSkill()
        await skill.run_async({
            "markdown_content": content,
            "output_path": str(out_path),
            "title": title,
            "author": author,
        })
        log.info(f"PDF erstellt (Playwright): {out_path.name}")
        return {"status": "success", "format": "pdf", "path": _rel(out_path), "filename": out_path.name}
    except Exception as e:
        log.error(f"create_pdf Fehler: {e}", exc_info=True)
        return {"status": "error", "message": str(e)}


# ── DOCX ──────────────────────────────────────────────────────────

@tool(
    name="create_docx",
    description=(
        "Erstellt ein Word-Dokument (.docx) aus Titel und Inhalt. "
        "Unterstützt Markdown-ähnliche Formatierung: '# ' = Überschrift 1, "
        "'## ' = Überschrift 2, '- ' = Aufzählung, normaler Text = Absatz. "
        "Gibt den Dateipfad zurück."
    ),
    parameters=[
        P("title",   "string", "Titel des Dokuments", required=True),
        P("content", "string", "Inhalt mit Markdown-Formatierung", required=True),
        P("author",  "string", "Autor (optional)", required=False),
    ],
    capabilities=["document", "docx"],
    category=C.DOCUMENT
)
async def create_docx(title: str, content: str, author: str = "Timus") -> dict:
    def _build():
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.core_properties.author = author
        doc.core_properties.title  = title

        # Titel
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Datum
        p = doc.add_paragraph(f"Erstellt: {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  {author}")
        p.runs[0].font.size  = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(120, 120, 120)
        doc.add_paragraph()

        for raw_line in content.splitlines():
            line = raw_line.strip()
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith(("- ", "* ")):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line == "":
                doc.add_paragraph()
            else:
                doc.add_paragraph(line)

        out_path = _safe_filename(title, "docx")
        doc.save(str(out_path))
        return out_path

    try:
        out_path = await asyncio.to_thread(_build)
        log.info(f"DOCX erstellt: {out_path.name}")
        return {"status": "success", "format": "docx", "path": _rel(out_path), "filename": out_path.name}
    except Exception as e:
        log.error(f"create_docx Fehler: {e}", exc_info=True)
        return {"status": "error", "message": str(e)}


# ── XLSX ──────────────────────────────────────────────────────────

@tool(
    name="create_xlsx",
    description=(
        "Erstellt eine Excel-Tabelle (.xlsx) aus Spaltenüberschriften und Datenzeilen. "
        "headers: Liste der Spaltenüberschriften. "
        "rows: Liste von Zeilen, jede Zeile ist eine Liste von Werten. "
        "Gibt den Dateipfad zurück."
    ),
    parameters=[
        P("title",   "string", "Titel / Dateiname der Tabelle", required=True),
        P("headers", "array",  "Spaltenüberschriften, z.B. ['Name', 'Wert', 'Datum']", required=True),
        P("rows",    "array",  "Datenzeilen, z.B. [['Max', 42, '2026-01-01'], ...]", required=True),
        P("sheet",   "string", "Name des Arbeitsblatts (Standard: Tabelle1)", required=False),
    ],
    capabilities=["document", "xlsx"],
    category=C.DOCUMENT
)
async def create_xlsx(title: str, headers: list, rows: list, sheet: str = "Tabelle1") -> dict:
    def _build():
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet[:31]

        # Kopfzeile formatieren
        header_fill = PatternFill("solid", fgColor="2E7D32")
        header_font = Font(bold=True, color="FFFFFF")
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=str(header))
            cell.fill   = header_fill
            cell.font   = header_font
            cell.alignment = Alignment(horizontal="center")

        # Daten eintragen
        for row_idx, row in enumerate(rows, 2):
            for col_idx, value in enumerate(row, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)

        # Spaltenbreite automatisch anpassen
        for col in ws.columns:
            max_len = max((len(str(cell.value or "")) for cell in col), default=10)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

        out_path = _safe_filename(title, "xlsx")
        wb.save(str(out_path))
        return out_path

    try:
        out_path = await asyncio.to_thread(_build)
        log.info(f"XLSX erstellt: {out_path.name}")
        return {"status": "success", "format": "xlsx", "path": _rel(out_path), "filename": out_path.name}
    except Exception as e:
        log.error(f"create_xlsx Fehler: {e}", exc_info=True)
        return {"status": "error", "message": str(e)}


# ── CSV ───────────────────────────────────────────────────────────

@tool(
    name="create_csv",
    description=(
        "Erstellt eine CSV-Datei aus Spaltenüberschriften und Datenzeilen. "
        "Einfacheres Format als XLSX, kann überall geöffnet werden."
    ),
    parameters=[
        P("title",   "string", "Dateiname (ohne Endung)", required=True),
        P("headers", "array",  "Spaltenüberschriften", required=True),
        P("rows",    "array",  "Datenzeilen", required=True),
    ],
    capabilities=["document", "csv"],
    category=C.DOCUMENT
)
async def create_csv(title: str, headers: list, rows: list) -> dict:
    def _build():
        out_path = _safe_filename(title, "csv")
        with open(out_path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            writer.writerows(rows)
        return out_path

    try:
        out_path = await asyncio.to_thread(_build)
        log.info(f"CSV erstellt: {out_path.name}")
        return {"status": "success", "format": "csv", "path": _rel(out_path), "filename": out_path.name}
    except Exception as e:
        log.error(f"create_csv Fehler: {e}", exc_info=True)
        return {"status": "error", "message": str(e)}


# ── TXT ───────────────────────────────────────────────────────────

@tool(
    name="create_txt",
    description="Erstellt eine einfache Textdatei (.txt) und speichert sie in results/.",
    parameters=[
        P("title",   "string", "Dateiname (ohne Endung)", required=True),
        P("content", "string", "Textinhalt", required=True),
    ],
    capabilities=["document", "txt"],
    category=C.DOCUMENT
)
async def create_txt(title: str, content: str) -> dict:
    def _build():
        out_path = _safe_filename(title, "txt")
        out_path.write_text(content, encoding="utf-8")
        return out_path

    try:
        out_path = await asyncio.to_thread(_build)
        log.info(f"TXT erstellt: {out_path.name}")
        return {"status": "success", "format": "txt", "path": _rel(out_path), "filename": out_path.name}
    except Exception as e:
        log.error(f"create_txt Fehler: {e}", exc_info=True)
        return {"status": "error", "message": str(e)}
